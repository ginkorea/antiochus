import PyPDF2
import docx
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
import requests
from io import BytesIO

class Book:
    """Represents the book with file path, URL, or extracted text."""

    def __init__(self, file_path_or_url):
        self.file_path_or_url = file_path_or_url
        self.text = None
        self.type = self.get_book_type()

    def get_book_type(self):
        """Determines if the input is a URL or a file, and identifies the file type."""
        if self.file_path_or_url.startswith("http"):
            return 'url'
        try:
            return self.file_path_or_url.split('.')[-1].lower()  # Determine type based on file extension
        except IndexError:
            return None

    def read(self):
        """Reads the book or webpage and extracts text based on the type."""
        if self.type == 'pdf':
            self.text = self.read_pdf()  # Local PDF file
        elif self.type == 'txt':
            self.text = self.read_txt()  # Local or remote text file
        elif self.type == 'epub':
            self.text = self.read_epub()  # Local or remote EPUB file
        elif self.type == 'docx':
            self.text = self.read_docx()  # Local or remote DOCX file
        elif self.type == 'url':
            self.text = self.read_url()  # URL pointing to webpage
        else:
            raise ValueError(f"Unsupported file type or URL: {self.type}")
        return self.text

    def read_pdf(self):
        """Reads text from a local PDF file."""
        text = ""
        try:
            with open(self.file_path_or_url, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() or ""  # Handle cases where text extraction may fail
        except Exception as e:
            raise ValueError(f"Error reading PDF: {e}")
        return text

    def read_remote_pdf(self):
        """Reads text from a remote PDF file (URL)."""
        text = ""
        try:
            response = requests.get(self.file_path_or_url)
            response.raise_for_status()  # Ensure the URL was successfully fetched
            pdf_reader = PyPDF2.PdfReader(BytesIO(response.content))  # Use BytesIO to handle the PDF content as a stream

            # Iterate through the pages and extract text
            for page in pdf_reader.pages:
                text += page.extract_text() or ""  # Handle cases where text extraction may fail
        except Exception as e:
            raise ValueError(f"Error reading remote PDF: {e}")
        return text

    def read_txt(self):
        """Reads text from a plain text file."""
        try:
            if self.type == 'url':
                response = requests.get(self.file_path_or_url)
                response.raise_for_status()
                return response.text
            else:
                with open(self.file_path_or_url, 'r', encoding='utf-8') as file:
                    return file.read()
        except Exception as e:
            raise ValueError(f"Error reading TXT: {e}")

    def read_epub(self):
        """Reads text from an EPUB file with extensive debugging."""
        try:
            if self.type == 'url':
                response = requests.get(self.file_path_or_url)
                response.raise_for_status()
                epub_content = BytesIO(response.content)
                book = epub.read_epub(epub_content)
            else:
                book = epub.read_epub(self.file_path_or_url)

            text = ""
            for item in book.get_items():
                # Debug: Print the content type of each item
                print(f"Item type: {type(item)} - Content type: {getattr(item, 'get_content_type', lambda: 'N/A')()}")

                # Check if the item has text content and extract it
                if hasattr(item, 'get_content_type'):
                    content_type = item.get_content_type()

                    # We're only interested in document types like xhtml, html, or plain text
                    if content_type in ['application/xhtml+xml', 'text/html', 'text/plain']:
                        soup = BeautifulSoup(item.get_body_content(), 'html.parser')
                        extracted_text = soup.get_text(separator=' ', strip=True)
                        text += extracted_text + "\n"

                        # Debug: Show the extracted text for this content item
                        print(f"Extracted Text (first 500 chars): {extracted_text[:500]}")

            return text

        except Exception as e:
            raise ValueError(f"Error reading EPUB: {e}")

    def read_docx(self):
        """Reads text from a DOCX file."""
        try:
            if self.type == 'url':
                response = requests.get(self.file_path_or_url)
                response.raise_for_status()
                doc_content = BytesIO(response.content)
                doc = docx.Document(doc_content)
            else:
                doc = docx.Document(self.file_path_or_url)

            text = "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            raise ValueError(f"Error reading DOCX: {e}")
        return text

    def read_url(self):
        """Reads and extracts text from a webpage."""
        try:
            response = requests.get(self.file_path_or_url)
            response.raise_for_status()  # Raise an error for bad status codes
            soup = BeautifulSoup(response.content, 'html.parser')
            text = soup.get_text(separator=' ', strip=True)  # Extract visible text
            return text
        except Exception as e:
            raise ValueError(f"Error reading URL: {e}")
